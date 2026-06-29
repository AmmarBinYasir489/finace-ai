"""Build the beginner-friendly VoiceTrack project documentation."""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "VoiceTrack_Project_Documentation.docx"

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
MUTED = RGBColor(0x66, 0x66, 0x66)
LIGHT_BLUE_FILL = "E8EEF5"
LIGHT_GRAY_FILL = "F4F6F9"


def set_cell_fill(cell, fill: str) -> None:
    """Apply a simple cell background color."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    """Set Word cell margins in DXA."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in {"top": top, "bottom": bottom, "start": start, "end": end}.items():
        tag = f"w:{edge}"
        element = tc_mar.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    """Set an exact cell width."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[int]) -> None:
    """Apply fixed DXA table geometry."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row) -> None:
    """Tell Word this row is a repeating table header."""
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]):
    """Create a compact, readable table."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_fill(cell, LIGHT_BLUE_FILL)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(header)
        run.bold = True
        run.font.color.rgb = DARK_BLUE
        run.font.size = Pt(10)
    set_repeat_table_header(table.rows[0])

    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            paragraph = cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(value)
            run.font.size = Pt(9.5)
    set_table_width(table, widths)
    doc.add_paragraph()
    return table


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    """Add a styled heading."""
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = BLUE if level < 3 else DARK_BLUE


def add_bullet(doc: Document, text: str) -> None:
    """Add one bullet item using Word's list style."""
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.add_run(text)


def add_number(doc: Document, text: str) -> None:
    """Add one numbered item using Word's list style."""
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.add_run(text)


def add_callout(doc: Document, title: str, body: str) -> None:
    """Add a one-cell callout for important beginner notes."""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    set_cell_fill(cell, LIGHT_GRAY_FILL)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    set_table_width(table, [9360])
    paragraph = cell.paragraphs[0]
    title_run = paragraph.add_run(title + ": ")
    title_run.bold = True
    title_run.font.color.rgb = DARK_BLUE
    body_run = paragraph.add_run(body)
    body_run.font.size = Pt(10)
    doc.add_paragraph()


def configure_document(doc: Document) -> None:
    """Apply compact_reference_guide document tokens."""
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("VoiceTrack Project Documentation")
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = MUTED


def add_cover(doc: Document) -> None:
    """Create the opening page."""
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("VoiceTrack Project Documentation")
    run.font.name = "Calibri"
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = BLUE

    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run("Beginner-friendly guide to changes, files, and app workflow")
    subtitle_run.font.size = Pt(13)
    subtitle_run.font.color.rgb = MUTED

    meta_rows = [
        ["Project", "VoiceTrack - Offline AI Voice Expense Tracker"],
        ["Purpose", "Track income and expenses by typing or speaking plain English."],
        ["Runtime", "Fully offline after Python packages and local models are installed."],
        ["Prepared", date.today().strftime("%d %B %Y")],
        ["Main run command", r".\.venv\Scripts\python.exe main.py"],
    ]
    add_table(doc, ["Item", "Details"], meta_rows, [2700, 6660])

    add_callout(
        doc,
        "Plain English summary",
        "VoiceTrack listens to or reads a sentence, extracts the transaction details, lets you confirm them, saves them in SQLite, and redraws the dashboard using Python calculations.",
    )
    doc.add_page_break()


def add_content(doc: Document) -> None:
    """Write the documentation body."""
    add_heading(doc, "1. What VoiceTrack Does", 1)
    doc.add_paragraph(
        "VoiceTrack is a desktop expense tracker for Windows. Instead of filling many form fields manually, the user can type or speak a sentence such as 'yesterday I purchased soap for 300'. The app turns that sentence into a structured transaction with amount, type, category, date, time, and description."
    )
    add_bullet(doc, "It works offline after setup; no paid API or cloud database is required.")
    add_bullet(doc, "Ollama handles language understanding only.")
    add_bullet(doc, "Python handles saving, totals, balances, filters, charts, and exports.")

    add_heading(doc, "2. Main Updates Made", 1)
    update_rows = [
        ["Setup", "Avoided PowerShell activation problems by documenting direct virtualenv Python commands.", "User can run the app without changing Windows execution policy."],
        ["Ollama", "Added timeout handling, fallback model support, and a local rule-based fallback.", "The app remains usable if the model is slow or still loading."],
        ["Voice input", "Installed and configured offline voice packages with sounddevice and Vosk.", "Microphone input works offline after the Vosk model is present."],
        ["Microphone UI", "Changed the mic into a start/stop toggle and renamed it Microphone.", "User can start listening, speak, stop, then process the captured text."],
        ["Date handling", "Added relative date parsing for yesterday, last week, last month, days ago, weeks ago, and weekdays.", "Transactions can be saved to the date the user mentioned, not only today's date."],
        ["Categories", "Improved fallback category detection for shopping items like soap, shampoo, detergent, and toothpaste.", "Common household purchases no longer default to Other as often."],
        ["History", "Added transaction date plus entered date/time in the transaction list.", "User can see both when the spending happened and when it was entered."],
        ["Dashboard", "Added total balance, monthly income, monthly expenses, savings rate, and week/month comparison charts.", "Dashboard gives more useful financial context."],
        ["Reports", "Added category spending cards and report charts.", "Reports are easier to scan category by category."],
        ["UI design", "Redesigned sidebar, dashboard, transactions, reports, logo, and light/dark mode styling based on reference images.", "The app looks more like a modern income and expense tracker."],
        ["Chart labels", "Moved crowded donut labels into a clean legend with amounts.", "Small category slices no longer overlap on the chart."],
        ["Tests", "Expanded tests for extraction, DB read/write, filters, date parsing, fallback parsing, and balance math.", "Important behavior can be checked quickly after changes."],
    ]
    add_table(doc, ["Area", "What changed", "Why it matters"], update_rows, [1500, 3960, 3900])

    add_heading(doc, "3. Project Files Explained", 1)
    file_rows = [
        ["main.py", "Starts the app. It imports the UI runner and opens the desktop window."],
        ["voicetrack/config.py", "Reads settings from .env, such as Ollama model names, database path, timeout, and Vosk model path."],
        ["voicetrack/db.py", "Creates SQLite tables, saves transactions, reads history, calculates totals, groups categories, and exports CSV."],
        ["voicetrack/extractor.py", "Active natural-language pipeline: Qwen extractor first, orchestrator validator second, fallback only if Ollama fails."],
        ["voicetrack/prompts.py", "Stores the extractor and orchestrator prompts used by the local LLM."],
        ["voicetrack/fallback.py", "Last-resort local parser used only when Ollama is unavailable or times out."],
        ["voicetrack/voice.py", "Captures microphone audio and converts speech to text using the offline Vosk model."],
        ["voicetrack/charts.py", "Draws dashboard and report charts with matplotlib."],
        ["voicetrack/ui.py", "Builds the desktop screens: dashboard, add entry, transactions, reports, charts, preview, delete, export, and theme toggle."],
        ["tests/test_extractor.py", "Tests the extractor/orchestrator flow with fake Ollama responses."],
        ["tests/test_fallback.py", "Tests the last-resort fallback parser."],
        ["tests/test_db.py", "Tests SQLite creation, insert, read, summary, category totals, and delete."],
        ["README.md", "Beginner setup and run instructions."],
        ["CONCEPTS.md", "Short explanation of SQLite, Ollama, CustomTkinter, matplotlib, JSON, and voice recognition."],
        ["requirements.txt", "Core packages needed for the desktop app."],
        ["requirements-voice.txt", "Extra packages needed for offline microphone input."],
        [".env / .env.example", "Local configuration for model names, timeout, database path, and voice model path."],
    ]
    add_table(doc, ["File", "What it does"], file_rows, [2500, 6860])

    add_heading(doc, "4. Working Flow of the App", 1)
    doc.add_paragraph("The app has two input paths: typed text and spoken voice. Both paths meet at the same processing pipeline.")
    for step in [
        "User opens the app with the run command.",
        "User goes to Add Entry.",
        "User types a sentence, or clicks Microphone, speaks, clicks Stop, and the recognized text appears in the input box.",
        "User clicks Process.",
        "The UI sends the text to extractor.extract().",
        "The first Qwen/Ollama call extracts transaction JSON.",
        "The orchestrator prompt receives the original sentence and that JSON, then corrects mistakes.",
        "If Ollama is unavailable or times out, fallback.py is used as a low-confidence safety net.",
        "The final transaction rows are saved into SQLite.",
        "The database saves the transaction in SQLite.",
        "The UI refreshes dashboard numbers, charts, and transaction lists.",
    ]:
        add_number(doc, step)

    add_callout(
        doc,
        "Important boundary",
        "The AI never calculates totals. The AI only extracts meaning from text. Python reads the database and performs all math, filtering, summaries, and charts.",
    )

    add_heading(doc, "5. Simple Data Flow Diagram", 1)
    diagram = doc.add_paragraph()
    diagram.style = "Normal"
    diagram.add_run(
        "Voice/Text input -> Ollama or local fallback -> JSON fields -> Python validation -> SQLite database -> Python totals and filters -> CustomTkinter screens and matplotlib charts"
    ).bold = True

    add_heading(doc, "6. Database in Simple Words", 1)
    doc.add_paragraph(
        "SQLite is a small database saved as one local file. VoiceTrack stores it by default at C:\\Users\\<you>\\VoiceTrack\\data.db. There is no separate database server."
    )
    schema_rows = [
        ["transactions.id", "Automatic number for each saved transaction."],
        ["transactions.type", "income or expense."],
        ["transactions.amount", "The money value."],
        ["transactions.category", "One fixed category such as Shopping, Utilities, Salary, or Other."],
        ["transactions.description", "Short readable text for the transaction."],
        ["transactions.date", "The date the transaction happened."],
        ["transactions.time", "The time saved for the transaction."],
        ["transactions.created_at", "The date and time when the entry was entered into the app."],
        ["categories.name", "The allowed category names."],
    ]
    add_table(doc, ["Database field", "Meaning"], schema_rows, [2700, 6660])

    add_heading(doc, "7. Screens Explained", 1)
    screen_rows = [
        ["Dashboard", "Shows total balance, monthly income, monthly expenses, savings rate, category donut chart, and spending comparison for today/week/month/all."],
        ["Add Entry", "Where the user types or records a transaction, processes it with AI/fallback parsing, previews fields, edits if needed, and confirms saving."],
        ["Transactions", "Shows saved history in a table. It supports search, type/category/date filters, and delete buttons."],
        ["Reports", "Shows category spending cards, category chart, monthly income vs expense chart, and CSV export."],
        ["Light/Dark mode", "Top-right switch changes the whole interface between dark and light themes."],
    ]
    add_table(doc, ["Screen", "Purpose"], screen_rows, [2200, 7160])

    add_heading(doc, "8. How to Run the Project", 1)
    doc.add_paragraph("Open PowerShell in the project folder and use these commands.")
    command_rows = [
        ["Create virtual environment", r"python -m venv .venv"],
        ["Install app packages", r".\.venv\Scripts\python.exe -m pip install -r requirements.txt"],
        ["Install voice packages", r".\.venv\Scripts\python.exe -m pip install -r requirements-voice.txt"],
        ["Start Ollama model", r"ollama run llama3.2:3b"],
        ["Run VoiceTrack", r".\.venv\Scripts\python.exe main.py"],
        ["Run tests", r".\.venv\Scripts\python.exe -m pytest tests -p no:cacheprovider"],
    ]
    add_table(doc, ["Task", "Command"], command_rows, [2700, 6660])

    add_heading(doc, "9. Common Problems and Fixes", 1)
    problem_rows = [
        ["Activate.ps1 is blocked", "Do not activate the venv. Use .\\.venv\\Scripts\\python.exe directly."],
        ["pip is not recognized", "Use python -m pip through the virtualenv Python command."],
        ["customtkinter missing", "Install requirements.txt inside the virtual environment."],
        ["Ollama timed out", "Make sure Ollama is running, use llama3.2:3b for faster extraction, or let fallback parse simple entries."],
        ["Microphone does not work", "Install requirements-voice.txt and confirm the Vosk model folder exists."],
        ["Wrong date", "Check whether the sentence includes a supported phrase such as yesterday, last week, 2 days ago, or a YYYY-MM-DD date."],
        ["Wrong category", "Use the preview form to edit category before Confirm. The app keeps fixed categories for clean reports."],
        ["Crowded chart labels", "Fixed by moving donut labels into a legend."],
    ]
    add_table(doc, ["Problem", "Fix"], problem_rows, [2700, 6660])

    add_heading(doc, "10. Test Coverage", 1)
    doc.add_paragraph(
        "The tests use a fake extractor, so they do not require Ollama or a microphone. They focus on the app's important logic."
    )
    test_items = [
        "Valid input becomes a database row and triggers a UI-style saved callback.",
        "Missing amount returns a clear error.",
        "Missing category becomes Other.",
        "Database write/read is verified.",
        "Daily, weekly, monthly, and all-time filters are checked.",
        "Income minus expenses balance is calculated by Python.",
        "Fallback parser handles number words like three thousand.",
        "Relative dates like yesterday and last week are normalized.",
        "Shopping items like soap are categorized better.",
    ]
    for item in test_items:
        add_bullet(doc, item)

    add_heading(doc, "11. What to Remember", 1)
    add_bullet(doc, "VoiceTrack is local-first: Ollama, Vosk, SQLite, Python, and the GUI all run on the laptop.")
    add_bullet(doc, "The model is a helper for language extraction, not a calculator.")
    add_bullet(doc, "Every saved transaction can be reviewed before saving.")
    add_bullet(doc, "The dashboard and reports come from SQLite data and Python math.")
    add_bullet(doc, "The safest run command is .\\.venv\\Scripts\\python.exe main.py.")

    add_heading(doc, "12. Future Improvements", 1)
    future_items = [
        "Add an edit button for existing transactions.",
        "Add budget limits per category.",
        "Add recurring transaction templates for rent, salary, or bills.",
        "Add a backup button for the SQLite database.",
        "Add more natural language date phrases if the user needs them.",
        "Package into a Windows .exe with PyInstaller.",
    ]
    for item in future_items:
        add_bullet(doc, item)


def main() -> None:
    """Build the DOCX file."""
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_content(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
